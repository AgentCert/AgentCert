from docx import Document


def main() -> None:
    path = "/mnt/d/AgentCert_Solution_Architecture_Document.docx"
    doc = Document(path)

    print(f"paragraphs={len(doc.paragraphs)}")
    for index, paragraph in enumerate(doc.paragraphs[:260]):
        text = paragraph.text.strip()
        if text:
            print(f"P{index}: {text}")

    print(f"tables={len(doc.tables)}")
    for table_index, table in enumerate(doc.tables[:20]):
        print(f"TABLE {table_index} rows={len(table.rows)} cols={len(table.columns)}")
        for row_index, row in enumerate(table.rows[:12]):
            values = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            print(f"R{row_index}: " + " | ".join(values))
        print("---")


if __name__ == "__main__":
    main()