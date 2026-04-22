from __future__ import annotations

from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph


DOC_PATH = Path("/mnt/d/AgentCert_Solution_Architecture_Document.docx")
BACKUP_PATH = Path("/mnt/d/AgentCert_Solution_Architecture_Document.backup.docx")
FALLBACK_DOC_PATH = Path("/mnt/d/Studies/AgentCert/AgentCert_Solution_Architecture_Document.updated.docx")


def replace_paragraph(doc: Document, old_text: str, new_text: str) -> bool:
    def normalize(value: str) -> str:
        return " ".join(value.replace("\n", " ").split())

    normalized_old = normalize(old_text)
    normalized_new = normalize(new_text)

    for paragraph in doc.paragraphs:
        if normalize(paragraph.text.strip()) == normalized_new:
            return True

    for paragraph in doc.paragraphs:
        current = paragraph.text.strip()
        normalized_current = normalize(current)
        if current == old_text or normalized_current == normalized_old or current.startswith(old_text):
            paragraph.text = new_text
            return True
    return False


def replace_paragraph_startswith(doc: Document, prefix: str, new_text: str) -> bool:
    if paragraph_exists(doc, new_text):
        return True

    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(prefix):
            paragraph.text = new_text
            return True
    return False


def paragraph_exists(doc: Document, text: str) -> bool:
    return any(paragraph.text.strip() == text for paragraph in doc.paragraphs)


def add_multiline_paragraph(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        if index < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)


def add_bullet_paragraphs(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def find_paragraph_by_prefix(doc: Document, prefixes: list[str]) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for prefix in prefixes:
            if text.startswith(prefix):
                return paragraph
    return None


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        new_paragraph.text = text
    if style:
        new_paragraph.style = style
    return new_paragraph


def insert_paragraph_after_table(table: Table, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    table._tbl.addnext(new_p)
    new_paragraph = Paragraph(new_p, table._parent)
    if text:
        new_paragraph.text = text
    if style:
        new_paragraph.style = style
    return new_paragraph


def insert_table_after(doc: Document, paragraph: Paragraph, rows: int, cols: int, style: str = "Table Grid") -> Table:
    table = doc.add_table(rows=rows, cols=cols)
    table.style = style
    paragraph._p.addnext(table._tbl)
    return table


def add_multiline_to_paragraph(paragraph: Paragraph, lines: list[str]) -> None:
    paragraph.clear()
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        if index < len(lines) - 1:
            run.add_break(WD_BREAK.LINE)


def add_section_diagram_and_datamodel(doc: Document) -> None:
    if not paragraph_exists(doc, "6.2.1 Detailed Azure Deployment Diagram"):
        architecture_anchor = find_paragraph_by_prefix(
            doc,
            ["6.2 High-Level Architecture (Logical)", "6.2 High‑Level Architecture (Logical)"],
        )
        if architecture_anchor is not None:
            section_heading = insert_paragraph_after(architecture_anchor, "6.2.1 Detailed Azure Deployment Diagram")
            intro = insert_paragraph_after(
                section_heading,
                "The deployment view below maps edge access, AKS namespaces, and managed Azure dependencies used by AgentCert in production.",
            )
            diagram_block = insert_paragraph_after(intro)
            add_multiline_to_paragraph(
                diagram_block,
                [
                    "Infosys Hosted Platform -> Azure Region",
                    "  |",
                    "  +-- Entry and Network",
                    "  |    +-- Users, operators, CI/CD pipelines",
                    "  |    +-- Azure DNS",
                    "  |    +-- Application Gateway + WAF",
                    "  |    +-- AGIC or NGINX Ingress",
                    "  |",
                    "  +-- AKS Cluster",
                    "  |    +-- litmus-chaos: frontend, auth server, GraphQL, registry APIs",
                    "  |    +-- litmus-exp: subscriber, event-tracker, workflow-controller, chaos-operator, chaos-exporter",
                    "  |    +-- agent namespaces: flash-agent, install-agent, install-app, target applications, agent-sidecar",
                    "  |    +-- shared integrations: Prometheus MCP, Kubernetes MCP, LiteLLM path",
                    "  |",
                    "  +-- Azure Managed Services",
                    "       +-- Azure Container Registry",
                    "       +-- Key Vault",
                    "       +-- MongoDB or Cosmos DB for MongoDB",
                    "       +-- Azure OpenAI",
                    "       +-- Langfuse",
                    "       +-- Azure Monitor and Storage",
                ],
            )
            topology_table = insert_table_after(doc, diagram_block, rows=1, cols=3)
            topology_header = topology_table.rows[0].cells
            topology_header[0].text = "Zone"
            topology_header[1].text = "Core Components"
            topology_header[2].text = "Purpose"
            topology_rows = [
                (
                    "Entry and Network",
                    "Azure DNS, Application Gateway + WAF, AGIC or NGINX ingress",
                    "Secure north-south traffic into the AKS hosted platform",
                ),
                (
                    "AKS Cluster",
                    "litmus-chaos, litmus-exp, agent namespaces, MCP integrations",
                    "Runs control-plane APIs, execution controllers, and benchmark workloads",
                ),
                (
                    "Azure Managed Services",
                    "ACR, Key Vault, MongoDB/Cosmos DB, Azure OpenAI, Langfuse, Monitor, Storage",
                    "Provides images, secrets, model access, persistent state, telemetry, and artifacts",
                ),
            ]
            for zone, components, purpose in topology_rows:
                row = topology_table.add_row().cells
                row[0].text = zone
                row[1].text = components
                row[2].text = purpose

    if not paragraph_exists(doc, "8.1.1 MongoDB Data Model"):
        data_anchor = find_paragraph_by_prefix(doc, ["8.1 Logical Data Categories"])
        if data_anchor is not None:
            model_heading = insert_paragraph_after(data_anchor, "8.1.1 MongoDB Data Model")
            model_intro = insert_paragraph_after(
                model_heading,
                "AgentCert persists platform state in MongoDB collections, while Langfuse stores deep LLM traces and scoring observations. The model below captures the primary entities and links.",
            )
            model_table = insert_table_after(doc, model_intro, rows=1, cols=4)
            model_header = model_table.rows[0].cells
            model_header[0].text = "Collection"
            model_header[1].text = "Primary Keys / IDs"
            model_header[2].text = "Important Fields"
            model_header[3].text = "Relationships"
            model_rows = [
                (
                    "project",
                    "_id",
                    "name, members[].user_id, members[].role, state",
                    "project 1:N environment, agentRegistry, chaosExperiments",
                ),
                (
                    "environment",
                    "environment_id, project_id",
                    "type, infra_ids[]",
                    "environment 1:N apps_registrations",
                ),
                (
                    "chaosInfrastructures",
                    "infra_id, project_id",
                    "name, infra_scope, status",
                    "chaosInfrastructures 1:N chaosExperiments and chaosExperimentRuns",
                ),
                (
                    "apps_registrations",
                    "appId, projectId, environmentId",
                    "name, namespace, method, status",
                    "Belongs to environment and project",
                ),
                (
                    "agentRegistry",
                    "agentId, projectId",
                    "vendor, capabilities[], endpoint, status, langfuseConfig",
                    "Belongs to project and participates in benchmark runs",
                ),
                (
                    "chaosExperiments",
                    "experiment_id, project_id, infra_id",
                    "experiment_type, revision[], recent_experiment_run_details[]",
                    "chaosExperiments 1:N chaosExperimentRuns",
                ),
                (
                    "chaosExperimentRuns",
                    "experiment_run_id, experiment_id, project_id, infra_id",
                    "phase, execution_data, resiliency_score, faults_*",
                    "Run evidence linked back to experiment and infrastructure",
                ),
                (
                    "Supporting collections",
                    "Varies by collection",
                    "chaosProbes, chaosHubs, imageRegistry, serverConfig, gitops, faultStudios, user",
                    "Reference catalogs, governance settings, probe metadata, and workflow composition",
                ),
            ]
            for collection, identifiers, fields, relations in model_rows:
                row = model_table.add_row().cells
                row[0].text = collection
                row[1].text = identifiers
                row[2].text = fields
                row[3].text = relations

            relation_notes = insert_paragraph_after_table(
                model_table,
                "Data model notes: core Litmus collections commonly use snake_case keys such as project_id and experiment_id, while newer app and agent registry collections use camelCase keys such as projectId, appId, and agentId.",
            )
            insert_paragraph_after(
                relation_notes,
                "For dashboard performance, recent run snapshots are denormalized in chaosExperiments.recent_experiment_run_details, while full execution evidence remains in chaosExperimentRuns and Langfuse.",
            )


def polish_language_and_flow(doc: Document) -> None:
    polish_replacements = [
        ("Chaos‑driven experiment of AI agent behaviour.", "The platform runs controlled chaos experiments to test AI agent behaviour."),
        ("Definition and execution of controlled failure experiments.", "Teams can define and run controlled failure experiments in a repeatable way."),
        ("External, non‑intrusive observability of agent behaviour.", "Agent behaviour is observed externally without changing core agent logic."),
        ("Objective scoring and benchmarking of agent outcomes.", "Outcomes are scored and benchmarked with objective, repeatable rules."),
        ("Certification evidence generation and audit trail.", "Each run produces certification evidence and a clear audit trail."),
        ("Agent development or training", "Building or training the agent itself is outside this scope."),
        ("Prompt or model fine tuning", "Prompt tuning or model fine-tuning is outside this scope."),
        ("Runtime intervention or overriding agent decisions", "The platform does not override agent decisions during runtime."),
        ("Business workflow orchestration outside certification", "General business workflow orchestration outside certification is out of scope."),
        ("Provide evidence‑based trust in AI agents before production", "Build evidence-based trust in AI agents before production rollout."),
        ("Detect reasoning failures early, not post‑deployment", "Find reasoning failures early, before they become production incidents."),
        ("Enable objective, repeatable agent comparison", "Enable fair and repeatable comparison across agents and versions."),
        ("Support governance, risk, and compliance reviews", "Support governance, risk, and compliance reviews with clear evidence."),
        ("Reduce reliance on anecdotal or manual validation", "Reduce reliance on anecdotal checks and manual validation."),
        ("Solution Components + Enhancements", "The architecture combines core platform components with production-focused enhancements."),
        ("Solution Flow", "The workflow moves from experiment definition to execution, observation, scoring, and review."),
        ("Layers:", "The architecture is organized into the following logical layers:"),
        ("Access Layer", "Access layer: UI, API entry, and operator workflows."),
        ("Control Plane", "Control plane: validation, orchestration, and governance."),
        ("Execution Plane", "Execution plane: chaos workflow execution and event processing."),
        ("Observer Plane", "Observer plane: traces, logs, metrics, and correlation."),
        ("Benchmarking Layer", "Benchmarking layer: scoring, verdict generation, and evidence."),
        ("Data Plane", "Data plane: persistent state, run history, and trace-linked records."),
        ("Customer Workloads", "Customer workloads: target applications and onboarded agent runtimes."),
        ("Responsibilities:", "Responsibilities in this layer:"),
        ("Interfaces:", "Primary interfaces:"),
        ("Key Components:", "Key components in this layer:"),
        ("GraphQL API", "GraphQL API service for experiment lifecycle, registry operations, and orchestration requests."),
        ("Design Note:\n This plane owns intent but never executes chaos directly.", "Design note: this plane defines intent and policy, but chaos is executed only in the execution plane."),
        ("Characteristics:", "Operational characteristics:"),
        ("Collected Signals:", "Signals collected in this layer:"),
        ("Traces", "Traces from OpenTelemetry and Langfuse, linked by run and experiment identifiers."),
        ("Typical Metrics:", "Typical benchmark metrics:"),
        ("Important:\n Scoring logic is implemented in code, not LLM judgment.", "Important: scoring logic is implemented in deterministic code. LLMs can assist analysis, but they do not calculate final scores."),
        ("Stored Artefacts:", "Stored artifacts:"),
        ("Append‑only execution history", "Append-only execution history to preserve run lineage."),
        ("Immutable evaluation records", "Immutable evaluation records for audit reliability."),
        ("Separation of raw telemetry and derived scores", "Clear separation of raw telemetry and derived scores."),
        ("Full lineage from experiment → score → verdict", "Maintain full lineage from experiment to score to final verdict."),
        ("Assumptions", "Assumptions used by this architecture:"),
        ("Target systems permit chaos safely", "Target systems allow controlled chaos experiments to run safely."),
        ("Constraints", "Constraints to keep the platform safe and auditable:"),
        ("Chaos limited to defined policies", "Chaos execution is limited to approved and defined policies."),
        ("Certification is evidence‑based only", "Certification decisions are based only on recorded evidence."),
    ]

    for old_text, new_text in polish_replacements:
        replace_paragraph(doc, old_text, new_text)

    bridge_anchor = find_paragraph_by_prefix(doc, ["6.1 Architectural Principles"])
    if bridge_anchor is not None and not paragraph_exists(
        doc,
        "This section first shows the high-level deployment view, then explains each architectural layer in operational order.",
    ):
        insert_paragraph_after(
            bridge_anchor,
            "This section first shows the high-level deployment view, then explains each architectural layer in operational order.",
        )

    data_bridge_anchor = find_paragraph_by_prefix(doc, ["8.2 Key Design Choices"])
    if data_bridge_anchor is not None and not paragraph_exists(
        doc,
        "These choices support the reliability, governance, and audit requirements described in the next sections.",
    ):
        insert_paragraph_after(
            data_bridge_anchor,
            "These choices support the reliability, governance, and audit requirements described in the next sections.",
        )


def verify_logical_order(doc: Document) -> bool:
    expected = [
        "1. Introduction",
        "2. Scope",
        "3. Business Context",
        "4. Core Components and Technology Stack",
        "5. Solution Overview",
        "6. Solution Architecture",
        "7. Process Flow",
        "8. Data Design",
        "9. Key Design Decisions",
        "10. Non‑Functional Considerations",
        "11. Security & Governance",
        "12. Assumptions & Constraints",
        "13. Summary",
    ]
    def find_index_by_prefix(prefix: str) -> int | None:
        for idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip().startswith(prefix):
                return idx
        return None

    indices: list[int] = []
    for title in expected:
        idx = find_index_by_prefix(title)
        if idx is None:
            return False
        indices.append(idx)
    return indices == sorted(indices)


def update_paragraphs(doc: Document) -> None:
    replacements = [
        (
            "This document describes the Solution Architecture for AgentCert, an AI Agent Certification platform built to evaluate and certify agent behaviour under controlled failure scenarios using cognitive chaos engineering.",
            "This document describes the solution architecture for AgentCert, an AI agent benchmarking and certification platform built on LitmusChaos, a Go-based GraphQL control plane, a React web UI, MongoDB-backed registries, Langfuse and OpenTelemetry observability, and Azure OpenAI-backed agent execution paths.",
        ),
        (
            "Support for bring‑your‑own agents (no code instrumentation).",
            "Unsupported custom agent onboarding paths that do not expose the required REST, Helm, or tracing contract are out of scope.",
        ),
        (
            "Python (control, execution, evaluators)",
            "Go 1.24 for control-plane services and Python 3.12 for agent-side runtimes, sidecars, and automation utilities.",
        ),
        (
            "Go / Python (chaos operators)",
            "Go services for GraphQL, auth, and event tracking; Python services for sidecar proxying and agent integrations.",
        ),
        (
            "React",
            "React 17 and TypeScript for the operator web UI.",
        ),
        (
            "Kubernetes (AKS or equivalent)",
            "AKS for production deployments and WSL + minikube + Docker driver for local development.",
        ),
        (
            "Container‑based deployment",
            "Container-based deployment with Helm charts, Kubernetes YAML manifests, and script-driven environment synchronization.",
        ),
        (
            "GraphQL (orchestrator API)",
            "GraphQL API, auth service, agent/app registries, and MongoDB-backed run-state services.",
        ),
        (
            "Chaos Engineering Framework (Litmus‑based)",
            "LitmusChaos, Argo Workflows, chaos-operator, subscriber, event-tracker, workflow-controller, and chaos-exporter.",
        ),
        (
            "OpenTelemetry",
            "OpenTelemetry, Langfuse, LiteLLM callbacks, and OTEL-to-Langfuse trace stitching.",
        ),
        (
            "Evaluation runtimes",
            "Custom benchmarking and verdict flows for TTD, TTR, remediation success, and trace-linked scoring.",
        ),
        (
            "Persistent data stores (telemetry, scores, audit)",
            "MongoDB for platform state and Langfuse for LLM traces, observations, and scores.",
        ),
        (
            "Traces, logs, metrics via OTel",
            "Traces, logs, and metrics via OTEL, Langfuse, Kubernetes signals, and Azure Monitor in production.",
        ),
        (
            "Platform‑level dashboards",
            "Platform dashboards across GraphQL, agents, chaos infrastructure, and evaluation results.",
        ),
        (
            "AgentCert is a framework‑agnostic, agent‑neutral certification platform that validates:",
            "AgentCert is an agent-neutral benchmarking platform that validates how onboarded agents detect failures, reason over cluster state, invoke remediation workflows, and produce traceable evidence under controlled chaos scenarios.",
        ),
        (
            "The platform separates chaos execution from judgement, ensuring unbiased evaluation.",
            "The platform separates chaos execution, telemetry capture, and scoring while still using lightweight sidecar and metadata contracts for trace correlation and custom agent onboarding.",
        ),
        (
            "(Insert AgentCert “layered architecture” PPT here)",
            "The logical architecture is layered around an access layer (UI and ingress), a Go control plane (GraphQL and auth), a Litmus execution plane, an observer plane (Langfuse, OTEL, LiteLLM, sidecar, MCP integrations), a benchmarking layer, and a MongoDB/Langfuse data plane.",
        ),
        (
            "Accept experiment definitions",
            "Accept experiment, agent, and target application definitions from the UI and APIs.",
        ),
        (
            "Support UI and automation workflows",
            "Support operator-driven workflows plus scripted onboarding, builds, and local deployment automation.",
        ),
        (
            "Web UI (guided operations)",
            "React web UI served through ingress and optionally Azure Application Gateway in production.",
        ),
        (
            "CLI / CI‑CD integration",
            "Shell and PowerShell automation, Helm-based onboarding, and CI/CD-driven deployment entrypoints.",
        ),
        (
            "Validate experiment definitions",
            "Validate experiment requests, registry metadata, and environment configuration.",
        ),
        (
            "Enforce governance and access",
            "Enforce authentication, RBAC, namespace scoping, and deployment-time configuration policy.",
        ),
        (
            "Orchestrate downstream workflows",
            "Orchestrate downstream Litmus workflows, Helm deployments, and trace-linked scoring flows.",
        ),
        (
            "GraphQL API",
            "GraphQL API and Go control-plane handlers.",
        ),
        (
            "Authentication & RBAC",
            "Authentication service, JWT issuance, and cluster RBAC bootstrapping.",
        ),
        (
            "Experiment registry",
            "Experiment and app registries stored in MongoDB.",
        ),
        (
            "Agent registry",
            "Agent registry plus Helm-driven install-agent and install-app orchestration.",
        ),
        (
            "Event publisher",
            "Event publisher and integration points into subscriber, event-tracker, and MCP services.",
        ),
        (
            "Execute chaos deterministically",
            "Execute repeatable Litmus chaos workflows against Kubernetes workloads and supporting services.",
        ),
        (
            "Apply faults to target environments",
            "Apply faults to registered applications and benchmark targets without changing core control-plane logic.",
        ),
        (
            "Workflow Controller",
            "Argo workflow-controller.",
        ),
        (
            "Chaos Operator",
            "Litmus chaos-operator and related CRDs/controllers.",
        ),
        (
            "Chaos Exporter",
            "Chaos exporter, subscriber, and event-tracker services.",
        ),
        (
            "Runs in isolated namespaces",
            "Runs primarily in litmus-chaos and litmus-exp namespaces with RBAC isolation.",
        ),
        (
            "Targets environment, not agent code",
            "Targets application and cluster behaviour; agents integrate via APIs, Helm onboarding, and sidecar metadata rather than intrusive runtime patching.",
        ),
        (
            "Fully policy‑driven",
            "Driven by manifests, registries, workflow policies, and environment-configured images.",
        ),
        (
            "Capture unbiased telemetry",
            "Capture platform, experiment, and LLM telemetry with correlated run keys and trace IDs.",
        ),
        (
            "Traces",
            "OpenTelemetry spans and Langfuse traces.",
        ),
        (
            "Logs",
            "Kubernetes logs, workflow logs, and agent/runtime logs.",
        ),
        (
            "Metrics",
            "Kubernetes metrics, Prometheus targets, and evaluation timing metrics.",
        ),
        (
            "Observer services",
            "Langfuse clients, OTEL exporters, and GraphQL observability handlers.",
        ),
        (
            "MCP hooks",
            "Kubernetes and Prometheus MCP services plus agent tool integrations.",
        ),
        (
            "OpenTelemetry pipelines",
            "OTEL pipelines bridged with Langfuse REST traces and LiteLLM callbacks.",
        ),
        (
            "Convert telemetry into objective scores",
            "Convert telemetry into objective run scores, verdicts, and audit evidence.",
        ),
        (
            "Persist experiment metadata",
            "Persist experiment definitions, agent registry data, and run metadata in MongoDB.",
        ),
        (
            "Store scores, traces, and results",
            "Store scores, verdict context, and LLM trace artifacts across MongoDB and Langfuse.",
        ),
        (
            "Enable replay and audits",
            "Enable replay, trace drill-down, and certification evidence audits.",
        ),
        (
            "Experiment definitions",
            "Experiment definitions, manifests, and deployment parameters.",
        ),
        (
            "Run history",
            "Run history, verdict history, and agent execution details.",
        ),
        (
            "Telemetry slices",
            "Trace spans, observations, logs, and metrics slices.",
        ),
        (
            "Evaluation outputs",
            "TTD, TTR, success, and scoring outputs.",
        ),
        (
            "(Insert AgentCert end‑to‑end flow diagram here)",
            "End-to-end flow: operator defines a benchmark, GraphQL persists and validates it, Litmus workflows and infra services execute the fault, the agent responds through registered interfaces and LiteLLM/Azure OpenAI paths, OTEL/Langfuse capture telemetry, and results are stored for audit and display.",
        ),
        (
            "Central authentication and RBAC",
            "Central authentication, JWT-based access, and Kubernetes RBAC bootstrapping for litmus-chaos and litmus-exp.",
        ),
        (
            "Role‑based experiment access",
            "Role-based experiment access with namespace-scoped execution identities and service accounts.",
        ),
        (
            "Immutable certification outcomes",
            "Run outcomes and score records are treated as append-only evidence for certification review.",
        ),
        (
            "No access to agent internals",
            "No deep access to proprietary agent internals is required; optional sidecar and metadata integration is used for trace correlation only.",
        ),
        (
            "Agents are externally observable",
            "Agents expose observable behaviour through APIs, traces, or supported onboarding contracts.",
        ),
        (
            "Telemetry pipelines are available",
            "Langfuse and OTEL paths are available or can be enabled for scoring-quality visibility.",
        ),
        (
            "No modification of agent logic",
            "Core agent remediation logic is not rewritten by the platform, though deployment-time sidecars, env vars, or metadata contracts may be used.",
        ),
        (
            "AgentCert provides a structured, auditable, and scalable architecture to certify AI agent behaviour under failure.\n By separating control, chaos, observation, and judgment, it enables enterprises to move from trust by assumption to trust by evidence.",
            "AgentCert provides a structured, auditable, and scalable architecture for benchmarking AI agents under chaos conditions. By combining a Go control plane, Litmus execution services, Langfuse and OTEL observability, LiteLLM and Azure OpenAI integrations, and MongoDB-backed run state, the platform moves teams from trust-by-assumption to trust-by-evidence.",
        ),
    ]

    for old_text, new_text in replacements:
        replace_paragraph(doc, old_text, new_text)


def update_tables(doc: Document) -> None:
    def set_cell_if_exists(table: Table, row_index: int, col_index: int, text: str) -> None:
        if row_index < len(table.rows) and col_index < len(table.columns):
            table.rows[row_index].cells[col_index].text = text

    revision_table = doc.tables[0]
    existing_revision = any(
        row.cells[0].text.strip() == "14‑04‑2026"
        and row.cells[1].text.strip() == "1.1"
        for row in revision_table.rows[1:]
    )
    if not existing_revision:
        new_row = revision_table.add_row().cells
        new_row[0].text = "14‑04‑2026"
        new_row[1].text = "1.1"
        new_row[2].text = "Expanded with detailed Azure deployment topology, namespace mapping, and operational guidance"

    existing_demo_revision = any(
        row.cells[0].text.strip() == "14‑04‑2026"
        and row.cells[1].text.strip() == "1.2"
        for row in revision_table.rows[1:]
    )
    if not existing_demo_revision:
        new_row = revision_table.add_row().cells
        new_row[0].text = "14‑04‑2026"
        new_row[1].text = "1.2"
        new_row[2].text = "Added certification context, cognitive-chaos methodology, metric framework, and mitigation guidance from the demo presentation"

    tech_table = doc.tables[1]
    set_cell_if_exists(tech_table, 12, 1, "Python 3.12")
    set_cell_if_exists(tech_table, 12, 2, "Used by the sidecar proxy and Python-based agent services; sidecar Dockerfile is based on python:3.12-slim")
    set_cell_if_exists(tech_table, 18, 1, "WSL Ubuntu, minikube, Docker driver")
    set_cell_if_exists(tech_table, 18, 2, "Current local workflow uses WSL Ubuntu plus minikube and Docker image loading scripts")
    set_cell_if_exists(tech_table, 20, 2, "build-all.sh orchestrates build-and-deploy flows, LiteLLM, install-agent, sidecar, and flash-agent builds")

    prod_table = doc.tables[2]
    set_cell_if_exists(prod_table, 8, 2, "subscriber, event-tracker, and supporting infra services")
    set_cell_if_exists(prod_table, 8, 3, "Watches infra events, feeds experiment state, and supports execution-plane coordination")
    set_cell_if_exists(prod_table, 9, 2, "LitmusChaos, Argo Workflows, chaos-operator, workflow-controller, chaos-exporter")
    set_cell_if_exists(prod_table, 18, 1, "Langfuse on AKS or managed external deployment")
    set_cell_if_exists(prod_table, 18, 3, "End-to-end agent and LLM observability with OTEL-linked Langfuse traces")

    principle_table = doc.tables[3]
    set_cell_if_exists(principle_table, 2, 1, "Supports custom agents through REST, Helm, and env-based onboarding without coupling the platform to one model vendor")
    set_cell_if_exists(principle_table, 4, 1, "Primary telemetry is external; optional sidecar and metadata integration is used only for trace correlation")

    metric_table = doc.tables[4]
    set_cell_if_exists(metric_table, 1, 1, "Time to Detect the injected or observed fault")
    set_cell_if_exists(metric_table, 2, 1, "Time to Recover or remediate after detection")
    set_cell_if_exists(metric_table, 3, 1, "Benchmark goal completion and remediation success")
    set_cell_if_exists(metric_table, 4, 1, "Decision quality, trace quality, and verdict correctness")

    flow_table = doc.tables[5]
    set_cell_if_exists(flow_table, 1, 1, "Experiment, target application, and agent are defined or selected")
    set_cell_if_exists(flow_table, 2, 1, "GraphQL control plane validates inputs and persists run metadata")
    set_cell_if_exists(flow_table, 3, 1, "Litmus workflows and infra services are triggered")
    set_cell_if_exists(flow_table, 4, 1, "Chaos fault is injected and the agent responds through registered interfaces")
    set_cell_if_exists(flow_table, 5, 1, "OTEL, Langfuse, logs, and metrics are captured with run correlation")
    set_cell_if_exists(flow_table, 6, 1, "Custom evaluators compute TTD, TTR, success, and verdict context")
    set_cell_if_exists(flow_table, 7, 1, "Results, traces, and audit evidence are persisted and shown in the UI")

    decision_table = doc.tables[7]
    set_cell_if_exists(decision_table, 3, 1, "Keeps benchmark execution repeatable and scoring independently auditable")
    set_cell_if_exists(decision_table, 5, 0, "Minimal agent-side integration")
    set_cell_if_exists(decision_table, 5, 1, "Allows trace stitching and onboarding without rewriting core agent remediation logic")

    nfr_table = doc.tables[8]
    set_cell_if_exists(nfr_table, 2, 1, "Controller retries, rollout validation, namespace bootstrapping, and recoverable script automation")
    set_cell_if_exists(nfr_table, 4, 1, "OTEL, Langfuse, Kubernetes logs, and Azure Monitor in production")


def append_detailed_sections(doc: Document) -> None:
    if paragraph_exists(doc, "Appendix A. Detailed Azure Deployment Diagram"):
        return

    doc.add_page_break()
    doc.add_paragraph("Appendix A. Detailed Azure Deployment Diagram", style="Heading 1")
    doc.add_paragraph(
        "The following deployment view expands the high-level production topology and shows the primary north-south traffic path, the workload split across namespaces, and the external services used for model access, secrets, image delivery, and observability."
    )
    add_multiline_paragraph(
        doc,
        [
            "Enterprise Users / Operators / CI-CD Pipelines",
            "  |",
            "Azure DNS",
            "  |",
            "Azure Application Gateway + WAF",
            "  |",
            "AGIC or NGINX Ingress Controller on AKS",
            "  |",
            "AKS Cluster",
            "  |-- Namespace: litmus-chaos",
            "  |    |-- litmusportal-frontend",
            "  |    |-- litmusportal-auth-server",
            "  |    `-- litmusportal-server (GraphQL / control plane)",
            "  |",
            "  |-- Namespace: litmus-exp",
            "  |    |-- subscriber",
            "  |    |-- event-tracker",
            "  |    |-- workflow-controller",
            "  |    |-- chaos-operator",
            "  |    |-- chaos-exporter",
            "  |    |-- kubernetes-mcp-server",
            "  |    `-- prometheus-mcp-server",
            "  |",
            "  `-- Agent / benchmark namespaces",
            "       |-- flash-agent deployments",
            "       |-- install-agent and install-app workloads",
            "       |-- target applications under test",
            "       `-- agent-sidecar proxy containers where enabled",
            "",
            "AKS integrations",
            "  |-- Azure Container Registry for images",
            "  |-- Azure Key Vault for secrets and certificates",
            "  |-- Azure OpenAI accessed through LiteLLM and agent runtimes",
            "  |-- MongoDB or Azure Cosmos DB for MongoDB for platform state",
            "  |-- Langfuse for LLM traces and evaluation observations",
            "  |-- Azure Monitor / Log Analytics for platform telemetry",
            "  `-- Azure Storage for exported artifacts, reports, and backups",
        ],
    )
    add_bullet_paragraphs(
        doc,
        [
            "North-south traffic enters through Azure DNS and Application Gateway, while east-west traffic remains inside the AKS virtual network and namespace boundaries.",
            "The Go control plane in litmus-chaos owns operator APIs, auth, experiment registration, app and agent registry data, and coordination with Litmus execution services.",
            "Litmus execution services remain isolated in litmus-exp so that fault injection, workflow execution, and execution-plane eventing do not share the same blast radius as the user-facing API tier.",
            "Agent workloads can be deployed in dedicated namespaces, which allows per-team network policies, service accounts, quota boundaries, and cleanup automation.",
            "Observability data fans out to OTEL, Langfuse, Azure Monitor, and Kubernetes-native logs so that every benchmark run can be reconstructed during review.",
        ],
    )

    doc.add_paragraph("Appendix B. Namespace and Workload Mapping", style="Heading 1")
    doc.add_paragraph(
        "The platform relies on a stable namespace split so that control-plane changes, experiment-plane operations, and agent-specific workloads can be versioned and scaled independently."
    )
    namespace_table = doc.add_table(rows=1, cols=4)
    namespace_table.style = "Table Grid"
    header = namespace_table.rows[0].cells
    header[0].text = "Namespace / Scope"
    header[1].text = "Primary Workloads"
    header[2].text = "Purpose"
    header[3].text = "Operational Notes"

    namespace_rows = [
        (
            "litmus-chaos",
            "Frontend, auth server, GraphQL server",
            "User-facing access layer and control-plane coordination",
            "Scale horizontally behind ingress; back with MongoDB connectivity, TLS termination, and RBAC-aware service accounts",
        ),
        (
            "litmus-exp",
            "subscriber, event-tracker, workflow-controller, chaos-operator, chaos-exporter",
            "Experiment execution, workflow orchestration, and event propagation",
            "Treat as execution plane; watch controller health, workflow backlog, and CRD availability",
        ),
        (
            "agent namespaces",
            "flash-agent, install-agent, install-app targets, sidecars",
            "Host benchmarked agents and application workloads under test",
            "Use namespace-per-team or namespace-per-benchmark patterns for cleanup and quota control",
        ),
        (
            "monitoring",
            "Prometheus, exporters, optional dashboards",
            "Metrics collection and MCP-backed tool access",
            "Expose only internal endpoints; retain metrics long enough to investigate failed benchmark runs",
        ),
        (
            "platform integrations",
            "LiteLLM, Langfuse, storage, secrets providers",
            "External dependencies for model access, traces, and persistent evidence",
            "Prefer private endpoints, managed identity, and secret rotation through Key Vault",
        ),
    ]
    for namespace_name, workloads, purpose, notes in namespace_rows:
        row = namespace_table.add_row().cells
        row[0].text = namespace_name
        row[1].text = workloads
        row[2].text = purpose
        row[3].text = notes

    doc.add_paragraph("Appendix C. Detailed Deployment and Runtime Flow", style="Heading 1")
    add_bullet_paragraphs(
        doc,
        [
            "Build and packaging: images for the control plane, sidecar, install-agent, flash-agent, and LiteLLM are built from the repository scripts and pushed to Azure Container Registry or loaded into local minikube during development.",
            "Configuration and secrets: Helm values, Kubernetes manifests, and runtime environment variables define image tags, registry URLs, Langfuse endpoints, Azure OpenAI credentials, MongoDB connectivity, and sidecar metadata such as agent name and agent role.",
            "Cluster bootstrap: namespaces, RBAC, service accounts, CRDs, and Litmus controllers must be present before rollout logic can safely restart or patch workloads, which is why the deployment scripts now include namespace and first-install guards.",
            "Application rollout: the frontend, auth server, and GraphQL server are deployed first so the control plane becomes available before experiment services, MCP services, or benchmark agents begin registering themselves.",
            "Execution flow: once an experiment is created, GraphQL persists the request, Litmus workflows execute the fault, the target agent reacts through its supported interfaces, and the result is evaluated against time-to-detect, time-to-recover, and remediation-success signals.",
            "Observability flow: OTEL spans, Langfuse traces, LiteLLM callbacks, Kubernetes logs, and event-tracker updates share run identifiers so investigators can correlate one benchmark from UI request to final verdict.",
            "Evidence retention: MongoDB stores platform state and benchmark metadata, while Langfuse and optional storage exports preserve detailed trace evidence for certification review and replay.",
        ],
    )

    doc.add_paragraph("Appendix D. Security and Operations Notes", style="Heading 1")
    add_bullet_paragraphs(
        doc,
        [
            "Place Application Gateway, AKS node pools, MongoDB endpoints, and observability backends inside approved VNets and private endpoint boundaries for production deployments.",
            "Use Key Vault-backed secret injection or sealed-secret patterns instead of embedding credentials in manifests or script files.",
            "Separate node pools for control-plane pods, execution-plane pods, and heavier benchmark workloads when agent or chaos runs have materially different CPU and memory profiles.",
            "Monitor controller health, ReplicaSet creation, rollout progress, and namespace-level quotas because deployment failures can originate from cluster control-plane issues rather than only from manifest errors.",
            "Retain backup and restore procedures for MongoDB state, Langfuse data, Helm values, and generated benchmark evidence so production certification records remain recoverable.",
            "Document supported onboarding contracts for third-party agents, including REST endpoints, Helm packaging expectations, required env vars, and optional sidecar tracing fields, so platform assumptions remain explicit.",
        ],
    )


def append_demo_insights(doc: Document) -> None:
    if paragraph_exists(doc, "Appendix E. Enterprise Trust Gap and Certification Context"):
        return

    doc.add_page_break()
    doc.add_paragraph("Appendix E. Enterprise Trust Gap and Certification Context", style="Heading 1")
    doc.add_paragraph(
        "The demo presentation adds important context beyond infrastructure design: enterprises do not only need a deployment topology for AgentCert, they need a repeatable answer to the trust gap created by autonomous AI agent behaviour in production."
    )
    add_bullet_paragraphs(
        doc,
        [
            "AI agents make autonomous decisions in production, yet those decisions are non-deterministic and difficult to reproduce after the fact.",
            "Organizations already certify people, software, and infrastructure, but they typically do not certify AI agents before those agents touch customers, SLAs, or regulated processes.",
            "AgentCert addresses this gap by replacing anecdotal trust with evidence-backed benchmark outcomes, certification scorecards, and repeatable audit trails.",
            "The platform goal is not only fault execution; it is pre-production confidence building for agent behaviour, tool selection, remediation quality, and safety posture.",
        ],
    )

    doc.add_paragraph("Appendix F. Cognitive Chaos Benchmarking Model", style="Heading 1")
    doc.add_paragraph(
        "The demo deck distinguishes traditional chaos engineering from AgentCert's cognitive-chaos model. This distinction should remain explicit in the architecture narrative because it explains why the platform includes evaluation, ground truth, and trace-linked certification outputs instead of fault injection alone."
    )
    comparison_table = doc.add_table(rows=1, cols=3)
    comparison_table.style = "Table Grid"
    comparison_header = comparison_table.rows[0].cells
    comparison_header[0].text = "Dimension"
    comparison_header[1].text = "Traditional Chaos Engineering"
    comparison_header[2].text = "AgentCert Cognitive Chaos"
    comparison_rows = [
        ("Primary subject", "System components such as pods, networks, or infrastructure", "Agent reasoning, decision quality, tool use, and remediation behaviour"),
        ("Core question", "Did the system survive the fault?", "Did the agent detect, interpret, and respond correctly under the fault?"),
        ("Loop structure", "Open-loop fault injection and observation", "Closed-loop benchmark flow with adaptive evaluation and post-run scoring"),
        ("Output", "Resilience signal for the target environment", "Certification evidence, scorecards, risk profile, and weakness analysis"),
    ]
    for dimension, traditional, cognitive in comparison_rows:
        row = comparison_table.add_row().cells
        row[0].text = dimension
        row[1].text = traditional
        row[2].text = cognitive

    doc.add_paragraph("Appendix G. Fault Families and Evaluation Metrics", style="Heading 1")
    doc.add_paragraph(
        "The demo deck also broadens the evaluation model. AgentCert is not limited to a single resiliency score; it measures agent behaviour across multiple fault families and metric categories that should be reflected in the formal architecture document."
    )
    add_bullet_paragraphs(
        doc,
        [
            "Fault families include network, compute, application, database, and security disruptions, implemented through LitmusChaos experiments and workflow orchestration.",
            "Detection and resolution metrics include TTD, TTM or MTTR-style remediation timing, false positive or false negative rate, and successful resolution rate.",
            "Reasoning-quality metrics include root-cause accuracy, explanation quality, decision quality, and consensus-based LLM judge outputs where qualitative review is needed.",
            "Behavioural metrics include tool-selection accuracy, plan adherence, collateral damage, rollback actions, exploration versus exploitation balance, and confidence signals.",
            "Security and efficiency metrics include PII handling, content-safety posture, authentication correctness, token usage, API call volume, and execution cost.",
        ],
    )
    metrics_table = doc.add_table(rows=1, cols=3)
    metrics_table.style = "Table Grid"
    metrics_header = metrics_table.rows[0].cells
    metrics_header[0].text = "Metric Family"
    metrics_header[1].text = "Representative Measures"
    metrics_header[2].text = "Why It Matters"
    metric_rows = [
        ("Detection and recovery", "TTD, TTR, detection accuracy, remediation success", "Confirms whether the agent can identify and recover from injected faults within acceptable windows"),
        ("Reasoning quality", "Root-cause accuracy, reasoning quality score, explanation clarity", "Assesses whether the agent reached the right conclusion for the right reason"),
        ("Action quality", "Tool-selection accuracy, argument accuracy, plan adherence, collateral damage", "Measures execution correctness and the cost of wrong actions"),
        ("Security and governance", "PII redaction, malicious prompt handling, auth correctness", "Ensures the agent behaves safely under stress and adversarial conditions"),
        ("Efficiency", "Token usage, API call count, latency, model cost", "Quantifies the operational footprint of the benchmarked agent"),
    ]
    for family, examples, rationale in metric_rows:
        row = metrics_table.add_row().cells
        row[0].text = family
        row[1].text = examples
        row[2].text = rationale

    doc.add_paragraph("Appendix H. Evaluation Method and Challenge Mitigations", style="Heading 1")
    add_bullet_paragraphs(
        doc,
        [
            "Faults are expected to run repeatedly, often around thirty executions per scenario, so certification outcomes are based on statistical aggregation rather than one-off anecdotes.",
            "Ground truth should not be exposed to the benchmarked agent; it remains an evaluator-side construct used for comparison against the ideal response trajectory.",
            "Qualitative review can use an LLM council pattern with multiple independent judges, while all arithmetic and score aggregation should remain deterministic and code-driven.",
            "Certification outputs should include a scorecard, trust or risk profile, weakness analysis, and full trace linkage in Langfuse for review and replay.",
        ],
    )
    mitigations_table = doc.add_table(rows=1, cols=2)
    mitigations_table.style = "Table Grid"
    mitigations_header = mitigations_table.rows[0].cells
    mitigations_header[0].text = "Challenge"
    mitigations_header[1].text = "Mitigation Strategy"
    mitigations_rows = [
        ("Non-determinism across agent runs", "Run each fault multiple times and use statistical aggregation rather than single-pass scoring"),
        ("Single-judge bias", "Use a council of independent models or evaluators for qualitative review"),
        ("Arithmetic mistakes in qualitative evaluators", "Keep all score calculations and aggregation in deterministic code paths"),
        ("Context-window pressure during long traces", "Batch traces and observations into controlled evaluation windows"),
        ("No standard agent workflow", "Benchmark observable outputs and supported interfaces rather than framework-specific internals"),
        ("Streaming or multi-fault complexity", "Sequence faults in a controlled orchestration model and isolate scenarios for reproducibility"),
    ]
    for challenge, mitigation in mitigations_rows:
        row = mitigations_table.add_row().cells
        row[0].text = challenge
        row[1].text = mitigation

    doc.add_paragraph("Appendix I. Certification Outputs and Adoption Path", style="Heading 1")
    add_bullet_paragraphs(
        doc,
        [
            "Near-term enterprise value includes evidence-based deployment approvals, version-to-version agent comparisons, and pre-production discovery of reasoning or remediation weaknesses.",
            "Certification artifacts should enable governance teams to define approval tiers, track risk profiles, and compare agents against shared benchmark scenarios.",
            "The long-term platform direction can support cross-team agent benchmarking, marketplace-style comparison, and standardized resilience or safe-action scores, provided the benchmark methodology stays auditable and reproducible.",
        ],
    )


def main() -> None:
    if not DOC_PATH.exists():
        raise FileNotFoundError(DOC_PATH)

    if not BACKUP_PATH.exists():
        copy2(DOC_PATH, BACKUP_PATH)

    doc = Document(str(DOC_PATH))
    update_paragraphs(doc)
    update_tables(doc)
    add_section_diagram_and_datamodel(doc)
    polish_language_and_flow(doc)
    append_detailed_sections(doc)
    append_demo_insights(doc)
    order_ok = verify_logical_order(doc)
    output_path = DOC_PATH
    try:
        doc.save(str(DOC_PATH))
    except PermissionError:
        output_path = FALLBACK_DOC_PATH
        doc.save(str(FALLBACK_DOC_PATH))
        print(f"Original document is locked; wrote updated copy to {FALLBACK_DOC_PATH}")
    else:
        print(f"Updated {DOC_PATH}")
    print(f"SectionOrderOK {order_ok}")
    print(f"Backup  {BACKUP_PATH}")
    print(f"Output  {output_path}")


if __name__ == "__main__":
    main()