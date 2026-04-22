from docx import Document

doc = Document('/mnt/d/AgentCert_Solution_Architecture_Document.docx')
print(f'Total paragraphs: {len(doc.paragraphs)}')
print(f'Total tables: {len(doc.tables)}')
print()

# Print all headings and their paragraph indices
print("=== HEADINGS / STRUCTURE ===")
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading') or p.style.name in ('Title', 'Subtitle'):
        print(f"P{i} [{p.style.name}]: {p.text[:100]}")

print()
print("=== ALL TABLES ===")
for i, t in enumerate(doc.tables):
    first_row = ' | '.join(c.text[:20] for c in t.rows[0].cells)
    print(f"TABLE {i}: {len(t.rows)} rows x {len(t.columns)} cols | headers: {first_row}")

print()
print("=== LOGICAL ORDER CHECK ===")
# Find all section headings and check numerical order
import re
sections = []
for i, p in enumerate(doc.paragraphs):
    m = re.match(r'^(\d+(?:\.\d+)*)\s', p.text)
    if m:
        sections.append((i, m.group(1), p.text[:80]))

prev_nums = None
for idx, (pi, sec, text) in enumerate(sections):
    parts = list(map(int, sec.split('.')))
    print(f"P{pi}: {sec} - {text[:60]}")
